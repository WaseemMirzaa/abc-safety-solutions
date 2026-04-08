#!/usr/bin/env python3
"""Generate client-facing Scope of Work DOCX for ABC Safety Solutions."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_table_from_rows(doc, headers, rows, header_fill="1a365d"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_fill)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title block ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Scope of Work")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(26, 54, 93)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Online Training Portal & Mobile Application")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(74, 85, 104)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "Prepared for: ABC Safety Solutions, Inc.",
        "Project: E‑commerce learning portal linked from existing website",
        "Document type: Scope & requirements summary",
    ):
        m = meta.add_run(line + "\n")
        m.font.size = Pt(11)
    doc.add_paragraph()

    hr = doc.add_paragraph("─" * 50)
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- 1. Purpose ---
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document defines the agreed scope for a dedicated online training platform "
        "that complements ABC Safety Solutions’ current marketing website. The existing "
        "WordPress site remains the public front door; a new tab or link will direct "
        "visitors to this portal for course discovery, purchase, learning, assessment, "
        "and digital credentials."
    )
    doc.add_paragraph(
        "The same capabilities will be available on iOS and Android so learners can "
        "register, complete training on mobile devices, and keep certificates readily "
        "available for job sites."
    )

    # --- 2. Objectives ---
    doc.add_heading("2. Business objectives", level=1)
    for item in (
        "Sell approximately twenty (20) online courses at launch, with room to add more "
        "through the admin panel without per‑course development fees for standard self‑service listings.",
        "Deliver each course as slide‑based content (typically 30–60 slides) with professional voice‑over audio.",
        "Require a knowledge check (test) after training; issue a branded PDF certificate upon successful completion.",
        "Support category-level certificate wording controlled by admin, so each training discipline can display its own certification text.",
        "Email the certificate automatically and store it in the learner’s web and app account.",
        "Support renewal reminders and optional broadcast announcements for new or updated offerings.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    # --- 3. In scope ---
    doc.add_heading("3. Scope of delivery", level=1)

    doc.add_heading("3.1 Web portal (new site)", level=2)
    web_items = (
        "Public pages: landing, course catalog, course detail, and checkout flow",
        "Branding: company logo, colors, and copy consistent with ABC Safety Solutions",
        "User accounts: sign‑up, login, password reset, profile (name used on certificates)",
        "E‑commerce: secure payment integration (e.g. Stripe) and order history",
        "Course player: slides with audio; progress saved so learners can resume",
        "Assessments: tests linked to courses; automated scoring; pass / fail handling",
        "Certificates: branded layout with prominent blue company logo; includes learner name, "
        "course title, completion date, and category-specific certification text; delivery by email and dashboard download",
        "Categories and subcategories for organizing the catalog",
        "Admin panel: manage courses, media, tests, pricing, users/orders, announcements, and category certificate text",
        "Responsive layout for phones and tablets",
        "Deployment-ready configuration for Vercel builds and Cloudflare Pages static build upload",
    )
    for x in web_items:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("3.2 Mobile applications", level=2)
    app_items = (
        "iOS and Android apps connected to the same backend as the web portal",
        "Browse or access purchased courses, complete training, and take tests",
        "Certificate wallet: view and present credentials on site",
        "Push notifications: training renewal / expiry reminders (where configured per course)",
        "Optional: broadcast messages to logged‑in users (e.g. new course announcements)",
    )
    for x in app_items:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("3.3 Integration with current website", level=2)
    doc.add_paragraph(
        "No redesign of the existing WordPress site is required. ABC Safety Solutions "
        "will add a navigation item (e.g. “Online courses”) that links to the new portal URL."
    )

    # --- 4. Agreed summary table ---
    doc.add_heading("4. Requirements summary (agreed)", level=1)
    add_table_from_rows(
        doc,
        ["Area", "Agreement"],
        [
            ("Website change", "Link/tab only to external portal; main site unchanged."),
            ("Initial catalog", "~20 courses; descriptions may mirror current Training & Certificates content."),
            ("Format", "Slide decks (~30–60 slides) + voice‑over audio."),
            ("Purchase", "Pay online → immediate access to purchased course(s)."),
            ("Completion", "Pass test → PDF certificate + email; stored in account."),
            ("Certificate content", "Admin can set/edit category-level certification text shown on certificates."),
            ("Branding on certificate", "Main blue web app logo is prominently displayed in certificate header."),
            ("Admin", "Full control to add/edit courses, tests, and media without developer for routine updates."),
            ("Mobile", "Same learning and certificates as web; reminders and announcements supported."),
            ("Deployment", "Configured for Vercel production builds and Cloudflare Pages `dist` uploads."),
        ],
    )

    # --- 5. User journeys ---
    doc.add_heading("5. Core user journeys", level=1)
    journeys = [
        ("Learner — web", "Main site → Portal → Choose course → Pay → Study slides/audio → Take test → Receive certificate by email → Download from dashboard."),
        ("Learner — app", "Install app → Sign in → My courses → Same learning and test flow → Certificates in wallet."),
        ("Administrator", "Log in → Create or edit course → Upload slides/audio → Build test → Set price and category → Publish → Monitor orders and completions."),
    ]
    for title_j, desc in journeys:
        p = doc.add_paragraph()
        p.add_run(title_j + ". ").bold = True
        p.add_run(desc)

    # --- 6. Milestones ---
    doc.add_heading("6. Phased delivery (indicative)", level=1)
    doc.add_paragraph(
        "Exact dates and payment splits will be confirmed in the project schedule. "
        "Typical phases:"
    )
    phases = [
        "Scope acceptance and environment planning (hosting, domains, app store accounts as needed)",
        "Design / front‑end prototypes (including mobile)",
        "Backend, admin panel, and APIs",
        "Payments, email, PDF certificates, and push notifications",
        "Content load and user acceptance testing (UAT) with ~20 courses",
        "Launch: portal live, WordPress link active, app store submission",
    ]
    for i, ph in enumerate(phases, 1):
        doc.add_paragraph(f"{i}. {ph}", style="List Number")

    doc.add_paragraph(
        "Note: Ongoing hosting, domain renewal, and third‑party fees (e.g. Apple/Google developer programs, "
        "email provider, cloud hosting) are typically borne by the client unless otherwise agreed."
    )

    # --- 7. Client responsibilities ---
    doc.add_heading("7. Client responsibilities", level=1)
    for c in (
        "Provide logo, brand guidelines, and final certificate PDF template (layout, signature, wording).",
        "Supply course copy, slide assets, and voice‑over audio (or approve production workflow).",
        "Provide list of ~20 launch courses and category structure.",
        "Decide payment processor and merchant‑of‑record approach; provide accounts as needed.",
        "Apple Developer and Google Play Console access for app publication (recommended under client account).",
        "Timely review and sign‑off at each milestone.",
    ):
        doc.add_paragraph(c, style="List Bullet")

    # --- 8. Items for joint confirmation ---
    doc.add_heading("8. Items to confirm (sign‑off)", level=1)
    doc.add_paragraph(
        "The following items were discussed and should be explicitly marked In scope or Phase 2 "
        "before build proceeds:"
    )
    add_table_from_rows(
        doc,
        ["#", "Topic", "Options / notes"],
        [
            ("1", "Payments", "Stripe on portal vs. alternate processor; refunds policy."),
            ("2", "App checkout", "In‑app purchase vs. web checkout + app access (store policy)."),
            ("3", "Certificate verification", "Public verify page / QR — yes, no, or later phase."),
            ("4", "Test rules", "Passing score, attempts allowed, time limit, question randomization."),
            ("5", "Slide / seat time gating", "None vs. must complete all slides before test."),
            ("6", "Accessibility", "Captions / transcripts for audio — phase 1 vs. 2."),
            ("7", "Hosting", "Provider choice and monthly budget."),
        ],
        header_fill="2c5282",
    )

    # --- 9. Out of scope ---
    doc.add_heading("9. Out of scope (unless added by change request)", level=1)
    for o in (
        "Redesign or migration of the entire existing WordPress marketing site",
        "Full enterprise LMS features (SCORM/xAPI, complex multi‑tenant portals)",
        "Live virtual classroom or webinar integration",
        "In‑person class scheduling and classroom management (separate from this online catalog)",
        "Legal compliance guarantees beyond implementing agreed features (client retains responsibility for training content accuracy and regulatory suitability)",
    ):
        doc.add_paragraph(o, style="List Bullet")

    # --- 10. Acceptance ---
    doc.add_heading("10. Acceptance", level=1)
    doc.add_paragraph(
        "Acceptance of this scope authorizes development to proceed according to the agreed "
        "milestones and pricing. Changes that materially expand functionality, integrations, "
        "or course count beyond self‑service admin use may be quoted separately."
    )
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Client approval\n\nName: _________________________\nTitle: ________________________\nDate: _________________________\n\n")
    sig.add_run("Developer / Vendor\n\nName: _________________________\nDate: _________________________\n")

    # Footer note
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        "This scope reflects discovery discussions with ABC Safety Solutions, Inc. "
        "and technical planning for the online training portal and mobile apps."
    ).italic = True
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(113, 128, 150)

    return doc


def main():
    out = Path(__file__).resolve().parent / "ABC_Safety_Solutions_Scope_of_Work.docx"
    doc = build_document()
    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
